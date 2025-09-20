#!/usr/bin/env python3

import sys
from docx import Document
import csv
import json

# Read the Word document
doc_path = '/Users/shivangpatel/Downloads/ICD The thar Dry Port (1).docx'
doc = Document(doc_path)

# Extract all text from the document
full_text = []
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        full_text.append(paragraph.text.strip())

# Also extract text from tables if any
for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            if cell.text.strip():
                row_text.append(cell.text.strip())
        if row_text:
            full_text.append(' | '.join(row_text))

# Print the content for analysis
print("=== DOCUMENT CONTENT ===\n")
for i, text in enumerate(full_text, 1):
    print(f"{i}. {text}")
    
print("\n=== END OF DOCUMENT ===")
print(f"\nTotal paragraphs/sections: {len(full_text)}")

# Save to a text file for easier processing
with open('/Users/shivangpatel/Documents/GitHub/crtx.in/docx_content.txt', 'w') as f:
    for text in full_text:
        f.write(text + '\n')
        
print("\nContent saved to docx_content.txt")